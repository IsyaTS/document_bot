from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from decimal import Decimal, InvalidOperation
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.company import Company
from app.parser import ParsedInput


DOC_TYPES = {
    "door_offer": "КП по дверям",
    "invoice": "Счет с QR",
    "measurement_estimate": "Смета по замеру",
    "door_contract": "Договор поставки и монтажа",
    "contract": "Договор",
    "measurement_act": "Акт замера",
    "installation_act": "Акт монтажных работ",
    "warranty": "Паспорт двери и гарантийный талон",
    "offer": "Коммерческое предложение",
    "claim_reply": "Ответ на претензию",
    "act": "Акт выполненных работ",
    "waybill": "Транспортная заявка",
    "reconciliation": "Акт сверки",
    "official_letter": "Исходящее письмо",
}

WARRANTY_MANUFACTURER = (
    "ООО «Феррони Йошкар-Ола», Республика Марий Эл, пгт. Медведево, "
    "ул. Железнодорожная, 11, www.ferroni-doors.ru"
)

WARRANTY_COVERAGE = [
    "Гарантийный срок эксплуатации дверного блока составляет 12 месяцев со дня отгрузки товара первому покупателю.",
    "Гарантийный срок эксплуатации замков и фурнитуры составляет 12 месяцев со дня отгрузки товара первому покупателю.",
    "Срок службы изделия при установке внутри помещений и соблюдении условий эксплуатации составляет 10 лет.",
    "Срок службы изделия при наружной установке по требованиям паспорта и соблюдении условий эксплуатации составляет 5 лет.",
    "Гарантийный ремонт выполняется при предъявлении документа, подтверждающего покупку у продавца.",
]

WARRANTY_MAINTENANCE = [
    "Закрывать и открывать замки и задвижку следует только после фиксации полотна на защелку.",
    "При заглубленной установке в проем необходимо предусмотреть ограничитель открывания или доводчик, иначе гарантия аннулируется.",
    "Во время влажных строительных процессов и при низких температурах возможно образование конденсата и обледенения, что не является гарантийным случаем.",
    "При появлении постороннего шума в петлевой зоне следует обратиться в сервисную службу; трущиеся части навесов подлежат смазыванию.",
    "Не реже одного раза в четыре месяца требуется смазывать защелку замка и доступные трущиеся поверхности запирающего механизма.",
]

WARRANTY_EXCLUSIONS = [
    "Механические повреждения вследствие неправильной транспортировки, хранения, монтажа или эксплуатации.",
    "Следы самостоятельного ремонта, разборки замков, фурнитуры или иного вмешательства в конструкцию.",
    "Повреждения после взлома или в результате обстоятельств непреодолимой силы.",
    "Установка доводчика и/или электромагнитного замка без предварительного согласования с производителем в спецификации на изделие.",
]

PASSPORT_GENERAL_INFO = [
    "Настоящий паспорт распространяется на двери производителя ООО «Феррони Йошкар-Ола», изготовленные по ТУ 25.12.10-005-03507917-2021, а также по спецификациям и конструкторской документации производителя.",
    "Изделие состоит из дверной коробки, дверного полотна и петель для навески полотна на коробку.",
    "Дверное полотно представляет собой сварную конструкцию из одного или двух стальных листов с усилениями и ребрами жесткости, кроме отдельных конструктивов типа Isoterma.",
    "Внутреннее заполнение может выполняться из пенополистирола, минераловатной плиты, базальтовой плиты или PIR-плиты в зависимости от спецификации.",
    "Дверной блок комплектуется замковыми устройствами, глазком, цилиндровыми механизмами, декоративными и/или броненакладками; состав определяется моделью и спецификацией.",
]

PASSPORT_TECH_SPECS = [
    "Двери могут изготавливаться с наружным или внутренним, правым или левым открыванием, а также с одной или двумя створками в зависимости от размеров блока.",
    "Габаритные размеры дверного блока могут изменяться по ширине от 700 до 1500 мм и по высоте от 1500 до 2200 мм.",
    "Внешний вид, отделка и комплектация соответствуют образцам и спецификациям производителя.",
    "В отделке могут использоваться декоративные панели, зеркальные элементы, молдинги, багеты, элементы из нержавеющей стали и иные опции по модели или спецзаказу.",
    "Производятся конструктивы 4,5 см, 6,0 см, 6,8 см, 7,5 см, 9 см, 10 см и Isoterma; конструктив Isoterma имеет терморазрыв и относится к теплосберегающим решениям.",
    "Двери без терморазрыва при установке с улицы в холодное время года могут промерзать или образовывать конденсат; такие явления не считаются дефектом.",
    "В блоках могут применяться механические и электронные замки, ночные задвижки, электронные ручки и дополнительные электронные системы согласно спецификации.",
]

PASSPORT_PACKAGE = [
    "Дверной блок - 1 шт.",
    "Паспорт - 1 шт.",
    "Ключи от замка(ов) - комплект, количество зависит от моделей замков.",
    "Комплект фурнитуры в упаковке - ручка с крепежом, вертушок ночной задвижки, цилиндровый механизм с ключами, глазок, эксцентрик, заглушки и иные элементы по модели.",
    "Комплект упаковки - 1 шт.",
]

PASSPORT_SAFETY = [
    "Изделия должны быть безопасными в эксплуатации и обслуживании, а применяемые материалы должны соответствовать действующим требованиям безопасности.",
    "К монтажу и техническому обслуживанию дверей допускается персонал с необходимой профессиональной подготовкой.",
]

PASSPORT_TRANSPORT = [
    "Транспортирование дверей в упаковке допускается любым видом наземного, воздушного и водного транспорта при обеспечении полной сохранности изделия.",
    "Кузов транспортного средства должен быть защищен от атмосферных осадков и посторонних предметов.",
    "При вертикальной перевозке двери располагаются с упором на жесткий передний борт; при горизонтальной перевозке не допускается укладка более 10 дверей друг на друга.",
    "Горизонтальная перевозка запрещена для изделий с ковкой и/или стеклянными элементами.",
    "Изделия следует хранить в вертикальном или горизонтальном положении на подкладках одинаковой толщины в закрытых вентилируемых помещениях, исключающих воздействие осадков и агрессивных сред.",
]

PASSPORT_OPERATION = [
    "Замки и задвижки следует открывать и закрывать только после постановки полотна на защелку.",
    "Люфт полотна на выдвинутых ригелях до 2 мм при снятом с защелки полотне не является дефектом.",
    "При заглубленной установке в проем обязательно применение ограничителя открывания или доводчика, иначе гарантия аннулируется.",
    "При нарушении температурно-влажностного режима в помещении возможно образование конденсата, инея и нарушение внешнего вида; восстановление в таком случае выполняется за счет покупателя.",
    "Двери со стеклопакетами не рекомендуется устанавливать как квартирные внутри подъезда из-за частичной прозрачности стеклопакета.",
    "При появлении скрипа или трения в петлевой части необходимо обратиться в сервисную службу; навесы подлежат смазке.",
    "Не реже одного раза в четыре месяца требуется смазывать защелку замка и трущиеся поверхности запирающего механизма.",
    "Профилактический осмотр и контроль работоспособности двери следует проводить не реже одного раза в три месяца.",
    "Металлические детали и декоративные панели рекомендуется протирать сухой или слегка влажной тканью со слабым мыльным раствором с последующим протиранием насухо.",
    "Во время строительно-отделочных работ следует защищать цилиндровый механизм от строительной пыли; после завершения работ производитель рекомендует замену цилиндра.",
]

PASSPORT_PROHIBITED = [
    "Запрещается устанавливать дверь с улицы без защитного козырька и холодного вентилируемого тамбура при необходимости их применения по условиям паспорта.",
    "Запрещается устанавливать в уличный проем отдельные модели с декоративными элементами из нержавеющей стали и панели, не предназначенные для наружной эксплуатации.",
    "Запрещается устанавливать доводчики и электромагниты на стандартные полотна Ferroni, изготовленные методом склеивания, без специального исполнения и согласования.",
    "Запрещается закрывать полотно при выдвинутых ригелях, при наличии посторонних предметов в зазоре или прилагать чрезмерное усилие к ключу.",
    "Запрещается самостоятельно разбирать и ремонтировать замок.",
    "Запрещается устанавливать не предусмотренные конструкцией дополнительные элементы декора, зеркала, замки, доводчики и прочее без отражения в спецификации.",
    "Запрещается воздействовать на покрытие абразивами, острыми предметами, химическими веществами и избыточной влагой.",
    "Запрещается превышать температурно-влажностный режим эксплуатации: температура от +5 °C до +40 °C, влажность от 35 % до 60 %.",
]

PASSPORT_INSTALLATION = [
    "Для монтажа используются: строительный уровень, рулетка, перфоратор с буром 10 мм, анкерные болты 10x120 мм, молоток, распорные клинья, отвертки, малярный скотч, монтажная пена и клей для наличников.",
    "Перед монтажом необходимо осмотреть дверь, проверить наружные зазоры и работоспособность замков и задвижки, убедиться в отсутствии механических повреждений.",
    "Проем подготавливается с монтажными зазорами не менее 10 мм по периметру; пол должен обеспечивать зазор под дверью не менее 5 мм.",
    "Короб выставляется по уровню и отвесу в двух плоскостях и закрепляется анкерами через технологические отверстия.",
    "После фиксации коробки проверяются зазоры между полотном и коробкой, работа замков, задвижки и сила прижима полотна к коробу.",
    "После окончательной проверки пространство между стеной и коробкой заполняется монтажной пеной, а после застывания ее излишки удаляются.",
]


@dataclass(frozen=True)
class GeneratedDocument:
    title: str
    path: Path
    filename: str


def generate_document(
    doc_type: str,
    company: Company,
    parsed: ParsedInput,
    output_dir: Path,
    ai_text: str | None = None,
) -> GeneratedDocument:
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _setup_styles(doc)

    title = DOC_TYPES.get(doc_type, "Документ")
    fields = parsed.fields
    number = fields.get("number") or date.today().strftime("%y%m%d")
    city = fields.get("city") or "Уфа"

    if doc_type == "door_offer":
        _build_door_offer(doc, company, parsed, number, city, ai_text)
    elif doc_type == "invoice":
        _build_invoice(doc, company, parsed, number, city)
    elif doc_type == "measurement_estimate":
        _build_measurement_estimate(doc, company, parsed, number, city, ai_text)
    elif doc_type == "door_contract":
        _build_door_contract(doc, company, parsed, number, city)
    elif doc_type == "contract":
        _build_contract(doc, company, parsed, number, city)
    elif doc_type == "measurement_act":
        _build_measurement_act(doc, company, parsed, number, city)
    elif doc_type == "installation_act":
        _build_installation_act(doc, company, parsed, number, city)
    elif doc_type == "warranty":
        _build_warranty(doc, company, parsed, number, city)
    elif doc_type == "offer":
        _build_offer(doc, company, parsed, number, city, ai_text)
    elif doc_type == "claim_reply":
        _build_claim_reply(doc, company, parsed, number, city, ai_text)
    elif doc_type == "act":
        _build_act(doc, company, parsed, number, city)
    elif doc_type == "waybill":
        _build_waybill(doc, company, parsed, number, city)
    elif doc_type == "reconciliation":
        _build_reconciliation(doc, company, parsed, number, city)
    else:
        _build_official_letter(doc, company, parsed, number, city, ai_text)

    slug = f"{doc_type}_{number}_{uuid4().hex[:8]}.docx"
    path = output_dir / slug
    doc.save(path)
    return GeneratedDocument(title=title, path=path, filename=slug)


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Pt(45)
        section.bottom_margin = Pt(45)
        section.left_margin = Pt(50)
        section.right_margin = Pt(35)


def _build_invoice(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Счет на оплату N {number} от {_today_ru()}")
    doc.add_paragraph(f"Поставщик: {company.data['full_name']}")
    _company_requisites_table(doc, company)
    doc.add_paragraph(f"Покупатель: {_counterparty_line(fields)}")
    if fields.get("basis"):
        doc.add_paragraph(f"Основание: {fields['basis']}")
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    doc.add_paragraph(f"Всего к оплате: {_format_money(_items_total(parsed.items))} руб.")
    doc.add_paragraph("Оплата данного счета означает согласие с условиями поставки/оказания услуг.")
    _signature(doc, company)


def _build_door_offer(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str, ai_text: str | None) -> None:
    fields = parsed.fields
    template = fields.get("template", "стандарт")
    _heading(doc, f"Коммерческое предложение по дверям N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Кому: {_counterparty_line(fields)}")
    _section(doc, "1. Предмет предложения")
    doc.add_paragraph(
        ai_text
        or (
            f"Предлагаем поставку дверей. Стандартный монтаж включен в стоимость изделия. Формат предложения: {template}. "
            "Итоговая комплектация и сроки фиксируются после подтверждения замера и наличия товара."
        )
    )
    _section(doc, "2. Параметры и комплектация")
    _door_specs(doc, fields)
    _section(doc, "3. Стоимость предложения")
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    _section(doc, "4. Условия")
    doc.add_paragraph(fields.get("payment_terms") or "Условия оплаты: предоплата по счету, окончательный расчет до передачи товара или в день монтажа.")
    doc.add_paragraph(fields.get("term") or "Срок поставки и монтажа согласуется после подтверждения заказа.")
    doc.add_paragraph("Если иное не согласовано отдельно, стоимость двери включает стандартный монтаж. Дополнительные работы и доставка указываются отдельно только при необходимости.")
    doc.add_paragraph(
        "Предложение носит коммерческий характер и действует для указанной комплектации. "
        "При изменении размеров, отделки, фурнитуры, способа монтажа или адреса доставки стоимость пересчитывается."
    )
    _signature(doc, company)


def _build_contract(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Договор N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(
        f"{company.data['full_name']}, именуемый в дальнейшем Исполнитель, в лице "
        f"{company.data['manager_name']}, с одной стороны, и {_counterparty_line(fields)}, именуемый "
        "в дальнейшем Заказчик, с другой стороны, заключили настоящий договор о нижеследующем."
    )
    _section(doc, "1. Предмет договора")
    doc.add_paragraph(fields.get("subject") or "Исполнитель обязуется оказать услуги/поставить товары по заявкам Заказчика, а Заказчик обязуется принять и оплатить результат.")
    _section(doc, "2. Цена и порядок расчетов")
    doc.add_paragraph(f"Стоимость работ/услуг составляет {_format_money(_declared_amount(parsed))} руб., {fields.get('vat') or company.data.get('tax_note', 'Без НДС')}.")
    doc.add_paragraph(fields.get("payment_terms") or "Оплата производится на расчетный счет Исполнителя на основании выставленного счета.")
    _section(doc, "3. Сроки исполнения")
    doc.add_paragraph(fields.get("term") or "Сроки исполнения согласуются сторонами в заявках, спецификациях или переписке.")
    _section(doc, "4. Ответственность сторон")
    doc.add_paragraph("Стороны несут ответственность за неисполнение обязательств в соответствии с законодательством Российской Федерации.")
    _section(doc, "5. Реквизиты и подписи сторон")
    _two_party_requisites(doc, company, fields)


def _build_door_contract(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Договор поставки и монтажа дверей N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(
        f"{company.data['full_name']}, именуемый в дальнейшем Исполнитель, в лице "
        f"{company.data['manager_name']}, с одной стороны, и {_counterparty_line(fields)}, именуемый "
        "в дальнейшем Заказчик, с другой стороны, заключили настоящий договор."
    )
    _section(doc, "1. Предмет договора")
    doc.add_paragraph(
        "Исполнитель обязуется изготовить и/или поставить дверной блок, комплектующие и иные согласованные позиции, а также при необходимости выполнить работы по доставке, демонтажу и монтажу, "
        "а Заказчик обязуется принять результат, обеспечить готовность объекта и оплатить заказ."
    )
    _door_specs(doc, fields)
    _section(doc, "2. Комплектация и стоимость")
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    _section(doc, "3. Замер, технические параметры и подготовка проема")
    _bullets(
        doc,
        [
            "размеры, технические условия, особенности стены, стороны открывания, отделка и иные параметры фиксируются в акте замера, спецификации или ином согласованном документе;",
            "Заказчик подтверждает корректность предоставленных размеров и обязан сообщить Исполнителю обо всех особенностях объекта, которые могут повлиять на монтаж;",
            "Заказчик обязан обеспечить готовность проема, свободный доступ к месту установки, возможность подъема изделия, наличие электропитания и отсутствие препятствий для проведения работ.",
        ],
    )
    _section(doc, "4. Права и обязанности Исполнителя")
    _bullets(
        doc,
        [
            "поставить изделие и выполнить согласованные работы в объеме, предусмотренном договором и согласованной спецификацией;",
            "использовать материалы и комплектующие, соответствующие модели, спецификации и конструктиву дверного блока;",
            "уведомлять Заказчика о выявленных обстоятельствах, требующих изменения состава работ, стоимости или сроков исполнения;",
            "отказать в выполнении монтажа до устранения препятствий, если объект не подготовлен или выполнение работ может повлечь риск повреждения изделия либо имущества.",
        ],
    )
    _section(doc, "5. Права и обязанности Заказчика")
    _bullets(
        doc,
        [
            "предоставить достоверные сведения о параметрах объекта и допустить представителей Исполнителя на объект в согласованное время;",
            "обеспечить сохранность свободной рабочей зоны, возможность транспортировки дверного блока до места монтажа и необходимые условия для проведения работ;",
            "принять изделие и работы, подписать документы приемки либо своевременно направить письменные мотивированные замечания;",
            "своевременно оплатить стоимость изделия, доставки, монтажа и дополнительных работ в согласованном объеме.",
        ],
    )
    _section(doc, "6. Сроки исполнения")
    doc.add_paragraph(fields.get("term") or "Срок поставки и выполнения работ согласуется сторонами после подтверждения заказа, комплектации, результатов замера и поступления предусмотренной договором оплаты.")
    doc.add_paragraph("Сроки подлежат соразмерному переносу при задержке оплаты, отсутствии готовности объекта, невозможности доступа на объект или при необходимости дополнительного согласования конструкции, отделки и комплектации.")
    _section(doc, "7. Порядок расчетов")
    doc.add_paragraph(fields.get("payment_terms") or "Оплата производится на основании счета Исполнителя. Размер аванса, промежуточных и окончательных платежей определяется согласованными условиями сделки.")
    doc.add_paragraph("Дополнительные работы, не включенные в исходную спецификацию, в том числе демонтаж, подъем, расширение проема, доборы, отделка откосов, доставка за пределы стандартного маршрута и иные услуги оплачиваются дополнительно после согласования с Заказчиком.")
    _section(doc, "8. Доставка, монтаж и дополнительные работы")
    _bullets(
        doc,
        [
            "стандартный монтаж включается в стоимость только в случаях, прямо согласованных сторонами; иные виды работ и услуг подлежат отдельному отражению в спецификации, счете или акте;",
            "монтаж выполняется после подтверждения готовности объекта и соблюдения технических требований к проему;",
            "при невозможности завершить работы по причинам, зависящим от Заказчика, Исполнитель вправе оформить частичное исполнение и перенести остаток работ на дополнительно согласованную дату.",
        ],
    )
    _section(doc, "9. Сдача и приемка")
    _bullets(
        doc,
        [
            "поставка изделия подтверждается накладной, универсальным передаточным документом либо иным документом передачи товара;",
            "выполнение монтажных и иных работ подтверждается актом выполненных работ или иным документом приемки;",
            "Заказчик обязан проверить комплектность, внешний вид, работоспособность замков, фурнитуры и отсутствие явных повреждений в момент приемки;",
            "после подписания документа приемки претензии по явным недостаткам, комплектности и внешнему виду принимаются только при наличии отметки в документе приемки.",
        ],
    )
    _section(doc, "10. Гарантия и правила эксплуатации")
    doc.add_paragraph("Гарантийные обязательства Исполнителя и производителя действуют при соблюдении правил эксплуатации, хранения, монтажа, технического обслуживания и иных требований паспорта изделия.")
    doc.add_paragraph("Гарантия не распространяется на случаи механического повреждения, неправильной транспортировки, самостоятельного ремонта, установки не предусмотренных конструкцией элементов, воздействия влаги, конденсата, химических веществ и иных нарушений условий эксплуатации.")
    _section(doc, "11. Ответственность сторон")
    doc.add_paragraph("Стороны несут ответственность за нарушение обязательств в соответствии с законодательством Российской Федерации. Заказчик несет риск последствий недостоверности предоставленных размеров и неподготовленности объекта.")
    doc.add_paragraph("Исполнитель не отвечает за скрытые дефекты проема, стен, инженерных сетей и иные обстоятельства объекта, которые не могли быть выявлены при обычном осмотре до начала работ.")
    _section(doc, "12. Форс-мажор")
    doc.add_paragraph("Стороны освобождаются от ответственности за неисполнение обязательств вследствие чрезвычайных и непредотвратимых обстоятельств, возникших после заключения договора.")
    _section(doc, "13. Разрешение споров и заключительные положения")
    doc.add_paragraph("Споры разрешаются путем переговоров, а при недостижении соглашения — в судебном порядке по месту нахождения Исполнителя, если иное не установлено обязательными нормами права.")
    doc.add_paragraph("Спецификации, счета, акты замера, акты выполненных работ, гарантийные документы и иная переписка сторон являются неотъемлемой частью настоящего договора.")
    _section(doc, "14. Реквизиты и подписи сторон")
    _two_party_requisites(doc, company, fields)


def _build_offer(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str, ai_text: str | None) -> None:
    fields = parsed.fields
    _heading(doc, f"Коммерческое предложение N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Кому: {_counterparty_line(fields)}")
    doc.add_paragraph(ai_text or _default_offer_text(fields))
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    doc.add_paragraph(fields.get("delivery") or "Сроки и условия доставки согласуются отдельно.")
    _signature(doc, company)


def _build_measurement_estimate(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str, ai_text: str | None) -> None:
    fields = parsed.fields
    _heading(doc, f"Смета по замеру N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Заказчик: {_counterparty_line(fields)}")
    doc.add_paragraph(f"Адрес объекта: {fields.get('object_address') or fields.get('counterparty_address', '-')}")
    _section(doc, "1. Основание составления сметы")
    doc.add_paragraph(
        ai_text
        or (
            "Смета составлена по результатам замера и на основании выбранной комплектации. "
            "Документ фиксирует расчет стоимости материалов, изделий и работ на текущую дату. "
            "Стандартный монтаж включается в стоимость двери, если отдельно не оговорено иное."
        )
    )
    rows = _door_specs_rows(fields)
    if rows:
        _section(doc, "2. Исходные данные по объекту")
        _key_value_table(doc, rows)
    _section(doc, "3. Сметный расчет")
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    _section(doc, "4. Примечания")
    doc.add_paragraph(
        "Смета не подтверждает факт выполнения работ и не заменяет договор или акт приемки. "
        "Итоговая стоимость может быть скорректирована после утверждения заказчиком состава поставки и работ."
    )
    _signature(doc, company)


def _build_claim_reply(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str, ai_text: str | None) -> None:
    fields = parsed.fields
    _heading(doc, f"Ответ на претензию N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"От: {company.data['full_name']}")
    doc.add_paragraph(f"Кому: {_counterparty_line(fields)}")
    doc.add_paragraph(ai_text or _default_claim_reply_text(fields))
    _signature(doc, company)


def _build_act(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Акт сдачи-приемки N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Исполнитель: {company.data['full_name']}")
    doc.add_paragraph(f"Заказчик: {_counterparty_line(fields)}")
    if fields.get("basis"):
        doc.add_paragraph(f"Основание: {fields['basis']}")
    _section(doc, "1. Состав выполненных работ / поставки")
    doc.add_paragraph("Стороны составили настоящий акт о том, что работы, услуги или поставка выполнены и переданы Заказчику в следующем объеме:")
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    _section(doc, "2. Подтверждение приемки")
    doc.add_paragraph("Заказчик принял результат работ и/или товар, комплектность проверена.")
    doc.add_paragraph("Претензий по объему, качеству и срокам оказания услуг стороны не имеют.")
    _two_party_signatures(doc, company, fields)


def _build_measurement_act(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Акт замера N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Исполнитель: {company.data['full_name']}")
    doc.add_paragraph(f"Заказчик: {_counterparty_line(fields)}")
    rows = [
        ("Адрес объекта", fields.get("object_address") or fields.get("counterparty_address", "")),
        ("Замерщик", fields.get("measurer", "")),
        ("Модель/тип двери", fields.get("model", fields.get("subject", ""))),
        ("Размер проема", fields.get("opening_size", "")),
        ("Толщина стены", fields.get("wall_depth", "")),
        ("Сторона открывания", fields.get("opening_side", "")),
        ("Цвет/отделка", fields.get("color", "")),
        ("Комментарий", fields.get("description", "")),
    ]
    _key_value_table(doc, rows)
    doc.add_paragraph(
        "Настоящий акт фиксирует результаты выезда и замера объекта. "
        "Документ не является коммерческим предложением, сметой или актом приемки монтажных работ."
    )
    doc.add_paragraph("Заказчик подтверждает корректность указанных размеров и технических условий на дату замера.")
    _two_party_signatures(doc, company, fields)


def _build_installation_act(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Акт выполненных монтажных работ N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Исполнитель: {company.data['full_name']}")
    doc.add_paragraph(f"Заказчик: {_counterparty_line(fields)}")
    doc.add_paragraph(f"Адрес объекта: {fields.get('object_address') or fields.get('counterparty_address', '-')}")
    if fields.get("basis"):
        doc.add_paragraph(f"Основание: {fields['basis']}")
    _section(doc, "1. Выполненные монтажные работы")
    doc.add_paragraph("Исполнитель выполнил поставку, монтаж дверей и сопутствующие работы в следующем составе:")
    _items_table(doc, parsed.items, fields.get("vat") or company.data.get("tax_note", "Без НДС"))
    _section(doc, "2. Результат приемки")
    doc.add_paragraph("Работы завершены, изделие установлено, объект передан Заказчику.")
    doc.add_paragraph("Претензий по объему и качеству на момент подписания акта не заявлено.")
    _two_party_signatures(doc, company, fields)


def _build_warranty(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Паспорт двери и гарантийный талон N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    _center(doc, "Входные двери. Паспорт на блок дверной стальной")
    doc.add_paragraph(
        "Документ оформлен на основании паспорта на блок дверной стальной производителя Ferroni "
        "и используется как гарантийный талон на реализованное изделие."
    )
    rows = [
        ("Производитель", WARRANTY_MANUFACTURER),
        ("Продавец/исполнитель", company.data["full_name"]),
        ("Покупатель", _counterparty_line(fields)),
        ("Адрес объекта", fields.get("object_address") or fields.get("counterparty_address", "")),
        ("Модель двери", fields.get("model", fields.get("subject", ""))),
        ("Размер", fields.get("size") or fields.get("opening_size", "")),
        ("Цвет/отделка", fields.get("color", "")),
        ("Серийный номер", fields.get("serial_number", "")),
        ("Дата продажи/установки", fields.get("term", _today_ru())),
        ("Срок гарантии", fields.get("warranty_period", "12 месяцев")),
    ]
    _key_value_table(doc, rows)
    _section(doc, "1. Общие сведения об изделии")
    for line in PASSPORT_GENERAL_INFO:
        doc.add_paragraph(f"- {line}")
    _section(doc, "2. Основные технические данные и характеристики")
    for line in PASSPORT_TECH_SPECS:
        doc.add_paragraph(f"- {line}")
    _section(doc, "3. Комплект поставки")
    for line in PASSPORT_PACKAGE:
        doc.add_paragraph(f"- {line}")
    _section(doc, "4. Требования безопасности")
    for line in PASSPORT_SAFETY:
        doc.add_paragraph(f"- {line}")
    _section(doc, "5. Транспортирование и хранение")
    for line in PASSPORT_TRANSPORT:
        doc.add_paragraph(f"- {line}")
    _section(doc, "6. Правила эксплуатации и технического обслуживания")
    for line in PASSPORT_OPERATION:
        doc.add_paragraph(f"- {line}")
    doc.add_paragraph("Запрещается:")
    for line in PASSPORT_PROHIBITED:
        doc.add_paragraph(f"- {line}")
    _section(doc, "7. Инструкция по монтажу")
    for line in PASSPORT_INSTALLATION:
        doc.add_paragraph(f"- {line}")
    _section(doc, "8. Гарантийные обязательства")
    for line in WARRANTY_COVERAGE:
        doc.add_paragraph(f"- {line}")
    _section(doc, "8.1. Условия эксплуатации и обслуживания")
    for line in WARRANTY_MAINTENANCE:
        doc.add_paragraph(f"- {line}")
    _section(doc, "8.2. Гарантийному ремонту не подлежит")
    for line in WARRANTY_EXCLUSIONS:
        doc.add_paragraph(f"- {line}")
    _section(doc, "9. Свидетельство о приемке и отметка о продаже")
    doc.add_paragraph("Адрес для претензий и замечаний: service@tk23.ru")
    doc.add_paragraph(
        "Блок дверной металлический соответствует требованиям паспорта и признан годным к эксплуатации "
        "при соблюдении правил монтажа и обслуживания."
    )
    doc.add_paragraph("Дата выпуска: ____________________    Контролер ОТК: ____________________")
    doc.add_paragraph("Дата установки: ____________________")
    _two_party_signatures(doc, company, fields)


def _build_waybill(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Транспортная заявка N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Исполнитель: {company.data['full_name']}")
    doc.add_paragraph(f"Заказчик: {_counterparty_line(fields)}")
    rows = [
        ("Маршрут", fields.get("route", fields.get("delivery", ""))),
        ("Адрес погрузки", fields.get("loading_address", "")),
        ("Адрес выгрузки", fields.get("unloading_address", "")),
        ("Груз", fields.get("subject", "")),
        ("Дата/срок", fields.get("term", "")),
        ("Стоимость", f"{_format_money(_declared_amount(parsed))} руб."),
        ("Условия оплаты", fields.get("payment_terms", "")),
    ]
    _key_value_table(doc, rows)
    _two_party_signatures(doc, company, fields)


def _build_reconciliation(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str) -> None:
    fields = parsed.fields
    _heading(doc, f"Акт сверки взаимных расчетов N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Между {company.data['full_name']} и {_counterparty_line(fields)} составлен настоящий акт сверки.")
    doc.add_paragraph(f"По данным сторон задолженность составляет {_format_money(_declared_amount(parsed))} руб.")
    doc.add_paragraph("При отсутствии письменных возражений в течение 5 рабочих дней акт считается согласованным.")
    _two_party_signatures(doc, company, fields)


def _build_official_letter(doc: Document, company: Company, parsed: ParsedInput, number: str, city: str, ai_text: str | None) -> None:
    fields = parsed.fields
    _heading(doc, f"Исходящее письмо N {number}")
    _center(doc, f"{city}, {_today_ru()}")
    doc.add_paragraph(f"Кому: {_counterparty_line(fields)}")
    doc.add_paragraph(ai_text or fields.get("description") or fields.get("subject") or "Просим рассмотреть настоящее письмо и сообщить о принятом решении.")
    _signature(doc, company)


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_heading(text, level=1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _center(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


def _bullets(doc: Document, rows: list[str]) -> None:
    for row in rows:
        doc.add_paragraph(f"- {row}")


def _company_requisites_table(doc: Document, company: Company) -> None:
    data = company.data
    rows = [
        ("ИНН/КПП", f"{data['inn']} / {data.get('kpp', '-')}"),
        ("Р/с", data["bank_account"]),
        ("Банк", data["bank_name"]),
        ("К/с", data["corr_account"]),
        ("БИК", data["bik"]),
    ]
    _key_value_table(doc, rows)


def _items_table(doc: Document, items: list[dict[str, str]], vat: str) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["N", "Наименование", "Кол-во", "Ед.", "Цена", "Сумма"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for idx, item in enumerate(items, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = item.get("name", "")
        row[2].text = item.get("qty", "1")
        row[3].text = item.get("unit", "усл.")
        row[4].text = _format_money(_decimal(item.get("price", "0")))
        row[5].text = _format_money(_decimal(item.get("total", "0")))
    doc.add_paragraph(f"Итого: {_format_money(_items_total(items))} руб. {vat}.")


def _key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value or "-"


def _door_specs(doc: Document, fields: dict[str, str]) -> None:
    rows = _door_specs_rows(fields)
    if rows:
        _section(doc, "Параметры заказа")
        _key_value_table(doc, rows)


def _door_specs_rows(fields: dict[str, str]) -> list[tuple[str, str]]:
    candidates = [
        ("Адрес объекта", fields.get("object_address") or fields.get("counterparty_address", "")),
        ("Модель/тип", fields.get("model", "")),
        ("Размер", fields.get("size") or fields.get("opening_size", "")),
        ("Толщина стены", fields.get("wall_depth", "")),
        ("Сторона открывания", fields.get("opening_side", "")),
        ("Цвет/отделка", fields.get("color", "")),
        ("Доставка", fields.get("delivery", "")),
        ("Монтаж", fields.get("installation", "")),
        ("Демонтаж", fields.get("dismantling", "")),
        ("Подъем", fields.get("lifting", "")),
        ("Доборы", fields.get("extras", "")),
        ("Наличники", fields.get("trim", "")),
        ("Фурнитура", fields.get("hardware", "")),
    ]
    return [(key, value) for key, value in candidates if value]


def _two_party_requisites(doc: Document, company: Company, fields: dict[str, str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    left, right = table.rows[0].cells
    left.text = _company_block(company)
    right.text = _counterparty_block(fields)


def _two_party_signatures(doc: Document, company: Company, fields: dict[str, str]) -> None:
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.text = f"{company.data['manager_title']}\n\n_____________ / {company.data['manager_name']} /"
    right.text = f"Заказчик\n\n_____________ / {fields.get('counterparty_manager', '')} /"


def _signature(doc: Document, company: Company) -> None:
    doc.add_paragraph("")
    doc.add_paragraph(f"{company.data['manager_title']} _______________ {company.data['manager_name']}")


def _company_block(company: Company) -> str:
    data = company.data
    return "\n".join(
        [
            data["full_name"],
            f"Адрес: {data['legal_address']}",
            f"ИНН/КПП: {data['inn']} / {data.get('kpp', '-')}",
            f"ОГРН/ОГРНИП: {data.get('ogrn', '-')}",
            f"Р/с: {data['bank_account']}",
            f"Банк: {data['bank_name']}",
            f"БИК: {data['bik']}",
            f"{data['manager_title']}: {data['manager_name']}",
        ]
    )


def _counterparty_block(fields: dict[str, str]) -> str:
    return "\n".join(
        [
            fields.get("counterparty_name", "Контрагент"),
            f"Адрес: {fields.get('counterparty_address', '-')}",
            f"ИНН/КПП: {fields.get('counterparty_inn', '-')} / {fields.get('counterparty_kpp', '-')}",
            f"ОГРН: {fields.get('counterparty_ogrn', '-')}",
            f"Телефон: {fields.get('counterparty_phone', '-')}",
            f"Email: {fields.get('counterparty_email', '-')}",
            f"Представитель: {fields.get('counterparty_manager', '-')}",
        ]
    )


def _counterparty_line(fields: dict[str, str]) -> str:
    name = fields.get("counterparty_name", "Контрагент")
    inn = fields.get("counterparty_inn")
    if inn:
        return f"{name}, ИНН {inn}"
    return name


def _default_offer_text(fields: dict[str, str]) -> str:
    subject = fields.get("subject") or "поставку товаров/оказание услуг"
    return f"Предлагаем рассмотреть условия сотрудничества по направлению: {subject}. Готовы согласовать объем, сроки и порядок оплаты под вашу заявку."


def _default_claim_reply_text(fields: dict[str, str]) -> str:
    claim = fields.get("claim_text") or "вашей претензии"
    return (
        f"Рассмотрев {claim}, сообщаем, что изложенные обстоятельства приняты в работу. "
        "По результатам проверки готовы предоставить дополнительные документы и согласовать порядок урегулирования."
    )


def _declared_amount(parsed: ParsedInput) -> Decimal:
    value = parsed.fields.get("amount")
    if value:
        return _decimal(value)
    return _items_total(parsed.items)


def _items_total(items: list[dict[str, str]]) -> Decimal:
    total = Decimal("0")
    for item in items:
        total += _decimal(item.get("total", "0"))
    return total


def _decimal(value: str) -> Decimal:
    try:
        clean = "".join(ch for ch in str(value).replace(",", ".") if ch.isdigit() or ch in ".-")
        return Decimal(clean or "0")
    except (InvalidOperation, ValueError):
        return Decimal("0")


def _format_money(value: Decimal) -> str:
    return f"{value.quantize(Decimal('0.01'))}"


def _today_ru() -> str:
    today = date.today()
    return today.strftime("%d.%m.%Y")
